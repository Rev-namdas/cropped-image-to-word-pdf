from glob import glob
from PIL import Image
import re
import docx2txt
import os
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import shutil

if os.path.isdir("img"):
    pass
else:
    os.mkdir("img")

file_to_process = "news_report.doc"

if os.path.isfile(file_to_process):
    os.rename(file_to_process, "news_report.docx")

file_to_process = "news_report.docx"

docx2txt.process(file_to_process, 'img/')

print("\nImage Extracted\n")

all_imgs = glob("img/*")
all_imgs.sort(key=lambda f: int(re.sub('\D', '', f)))

print("Image Grabbed\n")

if os.path.isdir("cropped"):
    pass
else:
    os.mkdir("cropped")

for each_img_url in all_imgs:
    each_img = Image.open(each_img_url)
    img_width, img_height = each_img.size
    how_many_times = round(img_height / img_width)

    if(how_many_times == 0):
        filename = each_img_url.split("\\")[1].split(".")[0]
        extension = each_img_url.split("\\")[1].split(".")[1]
        each_img.save("cropped/"+filename+"-"+str(each)+"."+extension)
    else:
        div_width = img_height / how_many_times
        crop_count = 0

        left = 1
        upper = 1
        right = img_width
        bottom = 1

        for each in range(how_many_times):
            if(crop_count == 0):
                upper = 1
                bottom = div_width
            elif((crop_count + 1) == how_many_times):
                upper = bottom
                bottom = img_height
            else:
                upper = bottom
                bottom = div_width * (each + 1)

            cropped_img = each_img.crop((left,upper,right,bottom))
            filename = each_img_url.split("\\")[1].split(".")[0]
            extension = each_img_url.split("\\")[1].split(".")[1]
            cropped_img.save("cropped/"+filename+"-"+str(each)+"."+extension)
            crop_count += 1

    print(f"{each_img_url} processed...")

document = Document()
header = document.sections[0].header
paragraph = header.paragraphs[0]
logo_runner = paragraph.add_run()
logo_runner.add_picture("header_logo.png", width=Inches(6))

footer = document.sections[0].footer
container = footer.paragraphs[0]
footer_runner = container.add_run()
footer_runner.add_picture("footer_logo.png", width=Inches(6))

all_imgs = glob("cropped/*")
all_imgs.sort(key=lambda f: int(re.sub('\D', '', f)))
count = 0

for each_img_url in all_imgs:
    count += 1
    document.add_picture(each_img_url, width=Inches(6), height=Inches(8))
    if(count != len(all_imgs)):
        print(f"Image {count} of {len(all_imgs)} Done")
        document.add_page_break()

document.save('img/AIS Report.docx')
convert("img/AIS Report.docx", "AIS Report.pdf")
print("\nYour Document Is Ready\n\n")

os.system("pause")
shutil.rmtree("img")
shutil.rmtree("cropped")


